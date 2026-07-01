"""
doc_generator.py
=================
Renders an ExecutionPlan + generated section content into a polished
.docx file using python-docx.
"""
from __future__ import annotations

import os
import re
from datetime import datetime
from typing import List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

from .models import ExecutionPlan, SectionContent

ACCENT_COLOR = RGBColor(0x2E, 0x5B, 0x8A)
MUTED_COLOR = RGBColor(0x66, 0x66, 0x66)


def _slugify(text: str) -> str:
    slug = re.sub(r"[^a-zA-Z0-9]+", "-", text.strip().lower()).strip("-")
    return slug or "document"


def _setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for level, size in ((1, 20), (2, 15)):
        style_name = f"Heading {level}"
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "Calibri"
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = ACCENT_COLOR


def _add_title_page(doc: Document, plan: ExecutionPlan) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(plan.document_title)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = ACCENT_COLOR

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(plan.document_type)
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = MUTED_COLOR

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_text = f"Audience: {plan.audience or 'N/A'}  |  Generated: {datetime.now().strftime('%B %d, %Y')}"
    meta_run = meta.add_run(meta_text)
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = MUTED_COLOR

    doc.add_paragraph()  # spacer


def _add_assumptions_box(doc: Document, assumptions: List[str]) -> None:
    if not assumptions:
        return
    heading = doc.add_heading("Agent Assumptions", level=2)
    note = doc.add_paragraph()
    note_run = note.add_run(
        "The following assumptions were made autonomously to resolve ambiguity or "
        "missing information in the original request:"
    )
    note_run.italic = True
    note_run.font.color.rgb = MUTED_COLOR
    for item in assumptions:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
    doc.add_paragraph()


def _add_section(doc: Document, section: SectionContent) -> None:
    doc.add_heading(section.title, level=2)
    for para in section.paragraphs:
        doc.add_paragraph(para)


def build_docx(
    plan: ExecutionPlan,
    sections: List[SectionContent],
    output_dir: str,
) -> str:
    os.makedirs(output_dir, exist_ok=True)
    doc = Document()

    # US Letter, 1" margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = section.right_margin = Inches(1)
    section.top_margin = section.bottom_margin = Inches(1)

    _setup_styles(doc)
    _add_title_page(doc, plan)
    _add_assumptions_box(doc, plan.assumptions)

    for sec in sections:
        _add_section(doc, sec)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{_slugify(plan.document_title)}_{timestamp}.docx"
    path = os.path.join(output_dir, filename)
    doc.save(path)
    return path
